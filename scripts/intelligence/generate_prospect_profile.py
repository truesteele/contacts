#!/usr/bin/env python3
"""
Generate formatted Prospect Profile .docx documents from Supabase contact data.

Usage:
    python generate_prospect_profile.py --id 3150          # by contact ID
    python generate_prospect_profile.py --name "Kaz McGrath"  # by name
    python generate_prospect_profile.py --ids 3150,255,2764   # multiple contacts
    python generate_prospect_profile.py --out /path/to/dir    # custom output dir

Output: docs/prospects/{First}_{Last}_Prospect_Profile.docx
"""

import argparse
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from dotenv import load_dotenv
from supabase import create_client

load_dotenv()

# --- Color Palette (matches existing examples) ---
GREEN_DARK = RGBColor(0x2D, 0x6A, 0x4F)   # #2D6A4F — accent, header bg
GREEN_LIGHT = "EAF4EE"                      # table row fill (alternating)
WHITE = "FFFFFF"
GRAY_LABEL = RGBColor(0x66, 0x66, 0x66)    # "Prospect Profile" label
TEXT_DARK = RGBColor(0x1A, 0x1A, 0x1A)      # name
FONT_NAME = "Arial"


def get_supabase():
    url = os.getenv("SUPABASE_URL") or os.getenv("NEXT_PUBLIC_SUPABASE_URL")
    key = os.getenv("SUPABASE_SERVICE_KEY") or os.getenv("SUPABASE_SERVICE_ROLE_KEY") or os.getenv("SUPABASE_KEY")
    if not url or not key:
        sys.exit("Missing SUPABASE_URL / SUPABASE_SERVICE_ROLE_KEY in .env")
    return create_client(url, key)


CONTACT_COLUMNS = """
    id, first_name, last_name, linkedin_url, headline, city, state, country,
    enrich_current_company, enrich_current_title, email,
    company, position,
    ai_tags, ask_readiness, communication_history,
    oc_engagement, fec_donations,
    comms_meeting_count, comms_last_meeting,
    comms_call_count, comms_last_call
""".strip()


def fetch_contact(sb, contact_id=None, name=None):
    """Fetch a single contact by ID or name."""
    q = sb.table("contacts").select(CONTACT_COLUMNS)
    if contact_id:
        q = q.eq("id", contact_id)
    elif name:
        parts = name.strip().split(None, 1)
        if len(parts) == 2:
            q = q.ilike("first_name", f"%{parts[0]}%").ilike("last_name", f"%{parts[1]}%")
        else:
            q = q.ilike("last_name", f"%{parts[0]}%")
    resp = q.execute()
    if not resp.data:
        return None
    return resp.data[0]


def fetch_threads(sb, contact_id):
    """Fetch all email/LinkedIn threads for a contact."""
    resp = (
        sb.table("contact_email_threads")
        .select("id, thread_id, account_email, subject, first_message_date, last_message_date, message_count, direction, channel, summary, raw_messages")
        .eq("contact_id", contact_id)
        .order("last_message_date")
        .execute()
    )
    return resp.data or []


# ─── Document helpers ───────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=Pt(11), color=None, font=FONT_NAME):
    """Set cell text with formatting, clearing existing content."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_kv_table(doc, rows_data):
    """Add a 2-column key-value table with alternating row shading."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set column widths
    for row_idx, (label, value) in enumerate(rows_data):
        row = table.rows[row_idx]
        fill = GREEN_LIGHT if row_idx % 2 == 0 else WHITE

        # Label cell
        set_cell_shading(row.cells[0], fill)
        set_cell_text(row.cells[0], label, bold=True, size=Pt(11))

        # Value cell
        set_cell_shading(row.cells[1], fill)
        set_cell_text(row.cells[1], value, bold=False, size=Pt(11))

    return table


def add_timeline_table(doc, headers, rows_data):
    """Add a multi-column timeline table with green header row."""
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for ci, header_text in enumerate(headers):
        cell = table.rows[0].cells[ci]
        set_cell_shading(cell, "2D6A4F")
        set_cell_text(cell, header_text, bold=True, size=Pt(10), color=RGBColor(0xFF, 0xFF, 0xFF))

    # Data rows
    for ri, row_data in enumerate(rows_data):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            # Only shade first column with alternating green
            if ci == 0:
                set_cell_shading(cell, GREEN_LIGHT if ri % 2 == 0 else WHITE)
            set_cell_text(cell, val, bold=False, size=Pt(10))

    return table


def add_heading(doc, text):
    """Add a section heading matching the template style."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = GREEN_DARK
    p.space_before = Pt(18)
    p.space_after = Pt(6)
    return p


def add_body(doc, text):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    p.space_after = Pt(6)
    return p


def add_mixed_line(doc, parts):
    """Add a line with alternating bold/normal segments. parts = [(text, bold), ...]"""
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.bold = bold
    p.space_after = Pt(4)
    return p


def add_callout(doc, text, icon=""):
    """Add a callout line (e.g. warning or success note)."""
    p = doc.add_paragraph()
    run = p.add_run(f"{icon}  {text}" if icon else text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.italic = True
    p.space_before = Pt(6)
    p.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    """Add a bullet-point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    # Clear default run and add formatted one
    p.clear()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    return p


# ─── Profile builder ───────────────────────────────────────────────

def build_profile(contact, threads, out_dir):
    """Build a Prospect Profile .docx for one contact."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    tags = contact.get("ai_tags") or {}
    ask = contact.get("ask_readiness") or {}
    comms = contact.get("communication_history") or {}
    oc = contact.get("oc_engagement") or {}
    fec = contact.get("fec_donations") or {}

    first = contact.get("first_name", "")
    last = contact.get("last_name", "")
    full_name = f"{first} {last}".strip()

    title = contact.get("enrich_current_title") or contact.get("position") or ""
    company = contact.get("enrich_current_company") or contact.get("company") or ""
    title_line = f"{title}, {company}" if title and company else title or company

    # Previous roles from tags
    prev_info = ""
    prox = tags.get("relationship_proximity", {})

    # ── HEADER ──
    p = doc.add_paragraph()
    run = p.add_run("Prospect Profile")
    run.font.name = FONT_NAME
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY_LABEL

    p = doc.add_paragraph()
    run = p.add_run(full_name)
    run.font.name = FONT_NAME
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = TEXT_DARK

    p = doc.add_paragraph()
    run = p.add_run(title_line)
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.color.rgb = GREEN_DARK

    doc.add_paragraph()  # spacer

    # ── PROFILE & CONTACT ──
    add_heading(doc, "Profile & Contact")

    profile_rows = []
    if title_line:
        profile_rows.append(("Role", title_line))
    loc_parts = [p for p in [contact.get("city"), contact.get("state"), contact.get("country")] if p]
    if loc_parts:
        profile_rows.append(("Location", ", ".join(loc_parts)))
    if contact.get("email"):
        profile_rows.append(("Email", contact["email"]))
    if contact.get("linkedin_url"):
        # Show clean URL
        url = contact["linkedin_url"].replace("https://www.", "").replace("https://", "")
        profile_rows.append(("LinkedIn", url))

    if profile_rows:
        add_kv_table(doc, profile_rows)

    # ── BACKGROUND ──
    add_heading(doc, "Background")

    # Build background from multiple tag sources
    bg_parts = []
    sf = tags.get("sales_fit", {})
    gc = tags.get("giving_capacity", {})
    ta = tags.get("topical_affinity", {})

    if sf.get("reasoning"):
        bg_parts.append(sf["reasoning"])
    if gc.get("reasoning") and gc["reasoning"] not in (sf.get("reasoning") or ""):
        bg_parts.append(gc["reasoning"])

    bg_text = " ".join(bg_parts) if bg_parts else contact.get("headline") or "No detailed background available."
    add_body(doc, bg_text)

    # ── RELATIONSHIP STATUS ──
    add_heading(doc, "Relationship Origin & Status")

    rel_rows = []
    if prox.get("tier"):
        score = prox.get("score", "")
        reasoning_short = prox.get("reasoning", "")[:120]
        rel_rows.append(("Proximity", f"{prox['tier'].title()} (score {score}) — {reasoning_short}"))

    # Familiarity from ask readiness
    ask_of = ask.get("outdoorithm_fundraising", {})
    if ask_of.get("reasoning"):
        # Extract familiarity mention
        reasoning = ask_of["reasoning"]
        fam_mention = ""
        if "familiarity" in reasoning.lower():
            for sentence in reasoning.split("."):
                if "familiarity" in sentence.lower():
                    fam_mention = sentence.strip()
                    break
        if fam_mention:
            rel_rows.append(("Familiarity", fam_mention))

    if prox.get("proximity_signals"):
        # Look for connection date
        for sig in prox["proximity_signals"]:
            if "connected" in sig.lower() or "linkedin" in sig.lower():
                rel_rows.append(("LinkedIn Connection", sig))
                break

    if prox.get("shared_employers") and len(prox["shared_employers"]) > 0:
        rel_rows.append(("Shared Employers", ", ".join(prox["shared_employers"])))
    if prox.get("shared_schools") and len(prox["shared_schools"]) > 0:
        rel_rows.append(("Shared Schools", ", ".join(prox["shared_schools"])))
    if prox.get("shared_boards") and len(prox["shared_boards"]) > 0:
        rel_rows.append(("Shared Boards", ", ".join(prox["shared_boards"])))

    if oc:
        roles = oc.get("crm_roles", [])
        donor = oc.get("is_oc_donor", False)
        trips = oc.get("trips_attended", 0)
        oc_summary = []
        if roles:
            oc_summary.append(f"CRM Role: {', '.join(roles)}")
        if donor:
            total = oc.get("oc_total_donated", 0)
            oc_summary.append(f"Donor (${total:,.0f} total)")
        if trips:
            oc_summary.append(f"{trips} trip(s) attended")
        if not donor and not trips:
            oc_summary.append("Not yet a donor or trip participant")
        if oc_summary:
            rel_rows.append(("OC Engagement", "; ".join(oc_summary)))

    if rel_rows:
        add_kv_table(doc, rel_rows)

    # ── COMMUNICATION TIMELINE ──
    add_heading(doc, "Communication Timeline")

    # Summary line
    total_threads = 0
    direction_counts = {"sent": 0, "received": 0, "bidirectional": 0, "outbound": 0, "inbound": 0}

    timeline_rows = []
    for t in threads:
        date_str = ""
        if t.get("last_message_date"):
            try:
                dt = datetime.fromisoformat(t["last_message_date"].replace("+00", "+00:00").split("+")[0])
                date_str = dt.strftime("%b %d, %Y")
            except:
                date_str = str(t["last_message_date"])[:10]

        channel = (t.get("channel") or "email").title()
        direction = t.get("direction") or "unknown"
        direction_display = direction.title()
        direction_counts[direction] = direction_counts.get(direction, 0) + 1

        summary = t.get("summary") or t.get("subject") or ""
        if len(summary) > 120:
            summary = summary[:117] + "..."

        timeline_rows.append((date_str, channel, direction_display, summary))
        total_threads += 1

    # Also include threads from communication_history that might not be in the threads table
    if comms and comms.get("threads") and not threads:
        for ct in comms["threads"]:
            date_str = ct.get("date", "")
            channel = "Email"
            direction = ct.get("direction", "unknown").title()
            summary = ct.get("summary", ct.get("subject", ""))
            if len(summary) > 120:
                summary = summary[:117] + "..."
            timeline_rows.append((date_str, channel, direction, summary))
            total_threads += 1

    # Stats line
    bidir = direction_counts.get("bidirectional", 0)
    replies = "Yes" if bidir > 0 else "None captured"
    meeting_count = contact.get("comms_meeting_count", 0) or 0
    call_count = contact.get("comms_call_count", 0) or 0
    stats_parts = [
        ("Email/DM threads: ", True), (f"{total_threads}   ", False),
        ("Bidirectional: ", True), (f"{replies}   ", False),
        ("Meetings: ", True), (f"{meeting_count}   ", False),
        ("Phone calls: ", True), (f"{call_count}   ", False),
        ("FEC Donations: ", True), (f"{'${:,.0f}'.format(fec.get('total_amount', 0)) if fec else '$0'}", False),
    ]
    add_mixed_line(doc, stats_parts)

    if timeline_rows:
        add_timeline_table(doc, ["Date", "Channel", "Direction", "Summary"], timeline_rows)
    else:
        add_body(doc, "No communication history on file.")

    # ── KEY RELATIONSHIP STRENGTHS (if enough data) ──
    strengths = []
    if bidir > 0:
        strengths.append("Bidirectional communication established — this is not a cold relationship.")
    if oc and oc.get("crm_roles"):
        strengths.append(f"Tracked in OC CRM as: {', '.join(oc['crm_roles'])}.")
    if prox.get("proximity_signals"):
        for sig in prox["proximity_signals"]:
            if "overlap" in sig.lower() or "shared" in sig.lower() or "both" in sig.lower():
                strengths.append(sig)
                break

    oc_ctx = tags.get("outreach_context", {})
    if oc_ctx.get("personalization_hooks"):
        for hook in oc_ctx["personalization_hooks"][:2]:
            strengths.append(hook)

    if strengths:
        add_heading(doc, "Key Relationship Strengths")
        for s in strengths:
            add_bullet(doc, s)
        doc.add_paragraph()  # spacer

    # ── FUNDRAISING ASSESSMENT ──
    add_heading(doc, "Fundraising Assessment")

    fund_rows = []
    if ask_of.get("tier"):
        tier_display = ask_of["tier"].replace("_", " ").title()
        score = ask_of.get("score", "")
        fund_rows.append(("Ask Readiness", f"{tier_display} (score {score})"))

    if gc.get("estimated_range"):
        fund_rows.append(("Personal Giving Capacity", gc["estimated_range"]))

    if ask_of.get("suggested_ask_range"):
        fund_rows.append(("Suggested Ask Range", ask_of["suggested_ask_range"]))

    if ask_of.get("recommended_approach"):
        fund_rows.append(("Recommended Channel", ask_of["recommended_approach"].replace("_", " ").title()))

    if sf.get("kindora_prospect") is not None:
        fit = "Yes" if sf["kindora_prospect"] else "No"
        if sf.get("prospect_type"):
            fit += f" — {sf['prospect_type'].replace('_', ' ').title()}"
        fund_rows.append(("Kindora Prospect", fit))

    if oc_ctx.get("outdoorithm_invite_fit"):
        fund_rows.append(("Outdoorithm Invite Fit", oc_ctx["outdoorithm_invite_fit"].title()))

    if fund_rows:
        add_kv_table(doc, fund_rows)

    # ── RECOMMENDED NEXT STEPS ──
    add_heading(doc, "Recommended Next Steps")

    # Callout
    if ask_of.get("tier") == "ready_now":
        add_callout(doc, "This contact is ready for an ask — proceed with a tailored outreach.", icon="\u2705")
    elif ask_of.get("tier") == "cultivate_first":
        if total_threads == 0 or (bidir == 0 and total_threads <= 3):
            add_callout(doc, "Key Risk: Another cold outreach without prior engagement will feel transactional. The next touch should be social/organic (LinkedIn) rather than a direct ask.", icon="\u26A0\uFE0F")
        else:
            add_callout(doc, "Relationship is warm enough to move toward an ask — but frame as a founding partner, not just a donor.", icon="\u2705")

    # Cultivation plan
    if ask_of.get("cultivation_needed"):
        # Split cultivation into bullet points
        cultivation = ask_of["cultivation_needed"]
        # Try to split on numbered items or semicolons
        steps = []
        if ":" in cultivation and ("(1)" in cultivation or "(2)" in cultivation):
            # Has numbered sub-steps — keep as one block but split on major sentences
            sentences = [s.strip() for s in cultivation.replace(". ", ".\n").split("\n") if s.strip()]
            for s in sentences:
                if len(s) > 15:
                    steps.append(s)
        else:
            steps = [cultivation]

        for step in steps[:6]:
            add_bullet(doc, step)
    elif ask_of.get("recommended_approach"):
        add_bullet(doc, f"Reach out via {ask_of['recommended_approach'].replace('_', ' ')}.")

    if ask_of.get("personalization_angle"):
        doc.add_paragraph()
        add_body(doc, f"Personalization angle: {ask_of['personalization_angle']}")

    # ── SAVE ──
    safe_name = f"{first}_{last}".replace(" ", "_").replace(".", "")
    filename = f"{safe_name}_Prospect_Profile.docx"
    filepath = os.path.join(out_dir, filename)
    os.makedirs(out_dir, exist_ok=True)
    doc.save(filepath)
    return filepath


# ─── Main ──────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate Prospect Profile documents")
    parser.add_argument("--id", type=int, help="Contact ID")
    parser.add_argument("--ids", help="Comma-separated contact IDs")
    parser.add_argument("--name", help="Contact name (first last)")
    parser.add_argument("--out", default="docs/prospects", help="Output directory")
    args = parser.parse_args()

    sb = get_supabase()

    contact_ids = []
    if args.id:
        contact_ids = [args.id]
    elif args.ids:
        contact_ids = [int(x.strip()) for x in args.ids.split(",")]
    elif args.name:
        c = fetch_contact(sb, name=args.name)
        if not c:
            sys.exit(f"No contact found matching '{args.name}'")
        contact_ids = [c["id"]]
    else:
        parser.print_help()
        sys.exit(1)

    for cid in contact_ids:
        contact = fetch_contact(sb, contact_id=cid)
        if not contact:
            print(f"  [SKIP] No contact with ID {cid}")
            continue

        threads = fetch_threads(sb, cid)
        filepath = build_profile(contact, threads, args.out)
        print(f"  [OK] {contact['first_name']} {contact['last_name']} → {filepath}")


if __name__ == "__main__":
    main()
